"""
Document Ingestors — Parse PDFs, DOCX, Markdown, and Code files.

These capture the twin owner's written thoughts, research, notes,
and code — forming the "knowledge base" component of the twin.
"""

import re
from datetime import datetime
from pathlib import Path
from typing import Optional
from loguru import logger

from . import BaseIngestor, Document, SourceType, DataCategory


class PDFIngestor(BaseIngestor):
    """Parse PDF documents into chunks."""
    
    def ingest(self) -> list[Document]:
        if not self.validate():
            return []
        
        documents = []
        
        for pdf_file in sorted(self.source_dir.glob("**/*.pdf")):
            try:
                # Use PyPDF2 for text extraction
                from PyPDF2 import PdfReader
                
                reader = PdfReader(str(pdf_file))
                full_text = ""
                
                for page_num, page in enumerate(reader.pages):
                    page_text = page.extract_text() or ""
                    full_text += page_text + "\n"
                
                if not full_text.strip():
                    logger.warning(f"No text extracted from {pdf_file}")
                    continue
                
                # Chunk the document
                chunks = self._chunk_text(full_text)
                
                for i, chunk in enumerate(chunks):
                    doc = Document(
                        content=chunk,
                        source_type=SourceType.PDF,
                        source_file=str(pdf_file),
                        doc_id=self._generate_doc_id("pdf", i, pdf_file.name),
                        category=DataCategory.GENERAL_KNOWLEDGE,
                        is_self=True,
                        metadata={
                            "filename": pdf_file.name,
                            "chunk_index": i,
                            "total_chunks": len(chunks),
                            "total_pages": len(reader.pages),
                        }
                    )
                    documents.append(doc)
                    
            except Exception as e:
                logger.warning(f"Error parsing PDF {pdf_file}: {e}")
        
        logger.info(f"PDF ingestor: parsed {len(documents)} chunks")
        self.documents = documents
        return documents
    
    def _chunk_text(self, text: str, chunk_size: int = 500, overlap: int = 50) -> list[str]:
        """Split text into overlapping chunks by words."""
        words = text.split()
        chunks = []
        start = 0
        
        while start < len(words):
            end = start + chunk_size
            chunk = " ".join(words[start:end])
            
            if chunk.strip() and len(chunk.strip()) > 50:
                chunks.append(chunk.strip())
            
            start = end - overlap
        
        return chunks


class DocxIngestor(BaseIngestor):
    """Parse Word documents."""
    
    def ingest(self) -> list[Document]:
        if not self.validate():
            return []
        
        documents = []
        
        for docx_file in sorted(self.source_dir.glob("**/*.docx")):
            try:
                from docx import Document as DocxDoc
                
                doc = DocxDoc(str(docx_file))
                paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                
                if not paragraphs:
                    continue
                
                # Group paragraphs into reasonable chunks
                chunks = self._group_paragraphs(paragraphs)
                
                for i, chunk in enumerate(chunks):
                    document = Document(
                        content=chunk,
                        source_type=SourceType.DOCX,
                        source_file=str(docx_file),
                        doc_id=self._generate_doc_id("docx", i, docx_file.name),
                        category=DataCategory.GENERAL_KNOWLEDGE,
                        is_self=True,
                        metadata={
                            "filename": docx_file.name,
                            "chunk_index": i,
                            "total_chunks": len(chunks),
                        }
                    )
                    documents.append(document)
                    
            except Exception as e:
                logger.warning(f"Error parsing DOCX {docx_file}: {e}")
        
        logger.info(f"DOCX ingestor: parsed {len(documents)} chunks")
        self.documents = documents
        return documents
    
    def _group_paragraphs(self, paragraphs: list[str], max_chars: int = 2000) -> list[str]:
        """Group paragraphs into chunks without exceeding max character count."""
        chunks = []
        current_chunk = []
        current_len = 0
        
        for para in paragraphs:
            if current_len + len(para) > max_chars and current_chunk:
                chunks.append("\n\n".join(current_chunk))
                current_chunk = []
                current_len = 0
            current_chunk.append(para)
            current_len += len(para)
        
        if current_chunk:
            chunks.append("\n\n".join(current_chunk))
        
        return chunks


class MarkdownIngestor(BaseIngestor):
    """Parse Markdown files (notes, journals, blog posts)."""
    
    def ingest(self) -> list[Document]:
        if not self.validate():
            return []
        
        documents = []
        
        for md_file in sorted(self.source_dir.glob("**/*.md")):
            try:
                content = md_file.read_text(encoding="utf-8")
                
                if not content.strip():
                    continue
                
                # Split by headers for natural chunking
                sections = self._split_by_headers(content)
                
                for i, section in enumerate(sections):
                    if len(section["content"].strip()) < 20:
                        continue
                    
                    doc = Document(
                        content=section["content"],
                        source_type=SourceType.MARKDOWN,
                        source_file=str(md_file),
                        doc_id=self._generate_doc_id("md", i, md_file.name),
                        category=self._classify_note(section["content"]),
                        is_self=True,
                        metadata={
                            "filename": md_file.name,
                            "section_header": section.get("header", ""),
                            "section_level": section.get("level", 0),
                        }
                    )
                    documents.append(doc)
                    
            except Exception as e:
                logger.warning(f"Error parsing Markdown {md_file}: {e}")
        
        logger.info(f"Markdown ingestor: parsed {len(documents)} sections")
        self.documents = documents
        return documents
    
    def _split_by_headers(self, content: str) -> list[dict]:
        """Split markdown content by headers."""
        sections = []
        current = {"header": "", "level": 0, "content": ""}
        
        for line in content.split("\n"):
            header_match = re.match(r"^(#{1,6})\s+(.+)", line)
            if header_match:
                if current["content"].strip():
                    sections.append(current)
                current = {
                    "header": header_match.group(2),
                    "level": len(header_match.group(1)),
                    "content": ""
                }
            else:
                current["content"] += line + "\n"
        
        if current["content"].strip():
            sections.append(current)
        
        return sections
    
    def _classify_note(self, content: str) -> DataCategory:
        """Basic classification of note content."""
        content_lower = content.lower()
        
        if any(w in content_lower for w in ["feel", "emotion", "afraid", "happy", "sad", "love"]):
            return DataCategory.DEEP_THOUGHTS
        if any(w in content_lower for w in ["meeting", "project", "deadline", "work", "team"]):
            return DataCategory.PROFESSIONAL
        if any(w in content_lower for w in ["journal", "diary", "today i", "reflection"]):
            return DataCategory.PERSONAL_STORIES
        if any(w in content_lower for w in ["decide", "choice", "should i", "option"]):
            return DataCategory.PRIVATE_DECISIONS
        
        return DataCategory.GENERAL_KNOWLEDGE


class CodeIngestor(BaseIngestor):
    """
    Parse code files — captures programming style and logic patterns.
    
    From the original 2020 system: "PROLOG became my everyday language" — 
    the twin needs to understand the owner's coding patterns too.
    """
    
    EXTENSIONS = {".py", ".pl", ".js", ".ts", ".html", ".css", ".java", ".go", ".rs"}
    
    def ingest(self) -> list[Document]:
        if not self.validate():
            return []
        
        documents = []
        
        for ext in self.EXTENSIONS:
            for code_file in sorted(self.source_dir.glob(f"**/*{ext}")):
                try:
                    content = code_file.read_text(encoding="utf-8", errors="ignore")
                    
                    if not content.strip() or len(content.strip()) < 20:
                        continue
                    
                    # Extract meaningful code sections (functions, classes, comments)
                    sections = self._extract_code_sections(content, ext)
                    
                    for i, section in enumerate(sections):
                        doc = Document(
                            content=section,
                            source_type=SourceType.CODE,
                            source_file=str(code_file),
                            doc_id=self._generate_doc_id("code", i, code_file.name),
                            category=DataCategory.PROFESSIONAL,
                            is_self=True,
                            metadata={
                                "filename": code_file.name,
                                "language": ext.lstrip("."),
                                "section_index": i,
                            }
                        )
                        documents.append(doc)
                        
                except Exception as e:
                    logger.warning(f"Error parsing code {code_file}: {e}")
        
        logger.info(f"Code ingestor: parsed {len(documents)} code sections")
        self.documents = documents
        return documents
    
    def _extract_code_sections(self, content: str, ext: str) -> list[str]:
        """Split code into logical sections."""
        lines = content.split("\n")
        sections = []
        current_section = []
        
        for line in lines:
            current_section.append(line)
            
            # Split on function/class definitions or after ~50 lines
            if len(current_section) >= 50:
                sections.append("\n".join(current_section))
                current_section = []
        
        if current_section:
            sections.append("\n".join(current_section))
        
        return [s for s in sections if s.strip()]
